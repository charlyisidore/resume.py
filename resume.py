#!/usr/bin/env python3

# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU General Public License as published by
# the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.

# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU General Public License for more details.

# You should have received a copy of the GNU General Public License
# along with this program.  If not, see <https://www.gnu.org/licenses/>.

# Author: Charly Lersteau

# ODT:
# ./resume.py examples/charly-lersteau.yml -f odt -t templates/resume.odt.yml.j2 -o resume.odt

# HTML:
# ./resume.py examples/charly-lersteau.yml -f html -t templates/resume.html.j2 -o resume.html

import argparse
import datetime
import dateutil.parser
import docx
import jinja2
import odf.opendocument
import odf.style
import odf.text
import odf.svg
import os
import pybtex.plugin
import pybtex.database
import tempfile
import yaml


# Parse command line arguments

parser = argparse.ArgumentParser(description="Resume generator")

parser.add_argument("input", type=str, help="YAML input file")

parser.add_argument(
    "-t", "--template", required=True, dest="template", type=str, help="Template file"
)

parser.add_argument(
    "-f", "--format", required=True, dest="format", type=str, help="Output format"
)

parser.add_argument(
    "-o", "--output FILE", required=True, dest="output", type=str, help="Output file"
)

args = parser.parse_args()

input_file = args.input
template_file = args.template
output_format = args.format
output_file = args.output


# Output functions


def text_output(output, file):
    with open(file, "w") as f:
        f.write(output)
        print("Output: {}".format(file))


def odf_output(output, file):
    def odf_load_yaml(data):
        doc = odf.opendocument.OpenDocumentText()
        yml = yaml.load(data, Loader=yaml.SafeLoader)
        for key, tree in yml.items():
            odf_load_tree(tree, getattr(doc, key))
        return doc

    def odf_load_tree(elements, parent):
        for el in elements:
            classname = el["class"].split(".")
            attrs = {}
            klass = odf
            while len(classname) > 0:
                klass = getattr(klass, classname[0])
                classname.pop(0)
            if "attrs" in el:
                attrs = el["attrs"]
            obj = klass(**attrs)
            if "children" in el:
                odf_load_tree(el["children"], obj)
            parent.addElement(obj)

    doc = odf_load_yaml(output)
    doc.save(file)
    print("Output odf: {}".format(file))


def docx_output(output, file):
    doc = docx.Document()
    doc.add_heading("Document Title", 0)
    p = doc.add_paragraph("A plain paragraph having some ")
    p.add_run("bold").bold = True
    p.add_run(" and some ")
    p.add_run("italic.").italic = True
    doc.add_page_break()
    doc.save(file)
    print("Output docx: {}".format(file))


output_function = {"html": text_output, "odt": odf_output, "docx": docx_output}


# Jinja2 template functions


def asstr(value):
    if value:
        if isinstance(value, list):
            return ", ".join(value)
        else:
            return value
    else:
        return ""


def aslist(value):
    if value != None:
        if not isinstance(value, list):
            return [value]
        else:
            return value
    else:
        return ""


def asdatetime(value, format):
    if isinstance(value, int):
        value = datetime.datetime(value, 1, 1)
    elif isinstance(value, str):
        value = dateutil.parser.parse(value)
    return value.strftime(format)


def asbib(value, exclude_fields=[]):
    if isinstance(value, str):
        path = os.path.join(os.path.dirname(input_file), value)
        bib = pybtex.database.parse_file(path, "bibtex")
    else:
        for field_name in exclude_fields:
            if field_name in value:
                del value[field_name]
        obj = {"entries": {"_": value}}
        yml = yaml.dump(obj, Dumper=yaml.SafeDumper)
        bib = pybtex.database.parse_string(yml, "yaml")
    fmt = pybtex_apa.format_bibliography(bib)
    return "".join(entry.text.render(pybtex_render) for entry in fmt)


# Pybtex configuration

pybtex_apa = pybtex.plugin.find_plugin("pybtex.style.formatting", "apa")()
if output_format == "html":
    pybtex_render = pybtex.plugin.find_plugin("pybtex.backends", "html")()
else:
    pybtex_render = pybtex.plugin.find_plugin("pybtex.backends", "plaintext")()


# Jinja2 configuration

file_system_loader = jinja2.FileSystemLoader("./")
environment = jinja2.Environment(loader=file_system_loader)

environment.filters["asstr"] = asstr
environment.filters["aslist"] = aslist
environment.filters["asdatetime"] = asdatetime
environment.filters["asbib"] = asbib

template = environment.get_template(template_file)


# Open the YAML input file

with open(input_file, "r") as f:
    data = yaml.load(f, Loader=yaml.SafeLoader)
    output = template.render(data)


# Output

if output_format in output_function:
    output_function[output_format](output, output_file)
else:
    print(output)
